"""
Generate RFP Word document (Request for Proposal) for the OGDCL document-grounded assistant.
Run: python scripts/generate_rfp_word.py
Output: docs/RFP_Request_for_Proposal_OGDCL_Document_Assistant.docx
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    if bold:
        run.bold = True


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    out_dir = root / "docs"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "RFP_Request_for_Proposal_OGDCL_Document_Assistant.docx"

    doc = Document()
    title = doc.add_heading("Request for Proposal (RFP)", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = "Calibri"

    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = st.add_run(
        "Intelligent Document Assistant for Corporate Research & Disclosure Analysis\n"
        "(OGDCL programme — reference corpus: U.S. SEC Form 10-K)"
    )
    r.font.name = "Calibri"
    r.font.size = Pt(12)
    r.bold = True

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m = meta.add_run(
        "\nIssuing function: Business & Digital Programme Office\n"
        "Document classification: Commercial / Technical — Request for Proposal\n"
        "Version: 1.0"
    )
    m.font.name = "Calibri"
    m.font.size = Pt(10)

    doc.add_page_break()

    # --- Executive summary ---
    add_heading(doc, "Executive Summary", 1)
    add_para(
        doc,
        "Oil & Gas Development Company Limited (OGDCL) seeks a vendor-implemented or internally "
        "delivered solution that enables authorised users to interrogate a large, structured regulatory "
        "filing—in this programme instantiated as The Coca-Cola Company Form 10-K (annual report)—through "
        "a natural-language interface, with auditable citations and strict grounding to source text. "
        "The capability supports market intelligence, peer benchmarking, and training use cases without "
        "replacing professional judgment or statutory reporting obligations.",
    )

    # --- PART A: Business Head ---
    doc.add_page_break()
    add_heading(doc, "Part A — Perspective of the Business Sponsor", 1)
    add_heading(doc, "1. Strategic intent & value proposition", 2)
    add_para(
        doc,
        "The organisation requires a defensible, scalable approach to extracting insight from lengthy "
        "public disclosures that would otherwise demand significant analyst time. The proposed solution "
        "must reduce time-to-insight, improve consistency of interpretation across teams, and create a "
        "repeatable pattern applicable to additional issuers or corpora without material re-engineering.",
    )
    add_heading(doc, "2. Commercial scope (high level)", 2)
    for line in [
        "Delivery of a production-grade web application and supporting API, or substantiated SaaS equivalent, aligned to OGDCL branding and access policy.",
        "Licensing and usage of third-party AI services (e.g. large language models and embeddings) must be transparent, auditable, and compliant with OGDCL procurement and data-governance standards.",
        "Total cost of ownership (TCO) clarity: implementation, hosting, support, token/API consumption, and change-management.",
        "Intellectual property: clear ownership of custom prompts, workflows, and deployment artefacts developed under this engagement.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    add_heading(doc, "3. Business outcomes & success measures", 2)
    for line in [
        "Measurable reduction in median time required to answer standard disclosure questions versus manual search (baseline to be agreed in mobilisation).",
        "User satisfaction score (e.g. ≥ 4/5) from pilot cohort after two weeks of live use.",
        "Zero critical incidents attributable to uncontrolled hallucination where the system asserts facts not present in the indexed corpus (to be validated by sampling).",
        "Availability target appropriate to internal research tooling (e.g. 99% monthly excluding planned maintenance), subject to chosen hosting tier.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    add_heading(doc, "4. Risk & compliance posture (business)", 2)
    add_para(
        doc,
        "Proponents shall address reputational, regulatory, and confidentiality risks: handling of API keys, "
        "segregation of production and non-production environments, retention of chat logs (if any), and "
        "alignment with OGDCL information security policies. Any cross-border data transfer inherent in "
        "cloud AI providers must be explicitly disclosed and mitigated.",
    )

    add_heading(doc, "5. Evaluation criteria (indicative weighting)", 2)
    add_para(
        doc,
        "Proposals will be scored on: (i) fit to stated outcomes, (ii) total cost and certainty of pricing, "
        "(iii) delivery timeline and governance model, (iv) quality of reference deployments, "
        "(v) strength of support and knowledge transfer. Weights may be finalised prior to vendor Q&A.",
    )

    # --- PART B: Technical Manager ---
    doc.add_page_break()
    add_heading(doc, "Part B — Perspective of the Technical Manager", 1)
    add_heading(doc, "1. Solution architecture (reference implementation)", 2)
    add_para(
        doc,
        "The reference architecture already implemented in the candidate solution comprises: "
        "(a) a FastAPI-based REST API exposing health, readiness, document ingestion, synchronous chat, "
        "and optional server-sent-event streaming; (b) LangChain-orchestrated retrieval-augmented generation "
        "(RAG) with hybrid semantic–lexical reranking; (c) ChromaDB as the vector persistence layer with "
        "OpenAI embedding models; (d) OpenAI chat completions for grounded answer synthesis; "
        "(e) a React (TypeScript) single-page application with optional 3D visualisation, built with Vite, "
        "consuming the API via configurable base URL for multi-environment deployment.",
    )

    add_heading(doc, "2. Functional requirements (technical)", 2)
    for line in [
        "Ingestion: support for large plain-text SEC-style filings and PDFs; configurable chunking (character targets, overlap), section-aware metadata (e.g. ITEM / PART boundaries, line ranges), and optional tail exclusion patterns.",
        "Indexing: idempotent rebuild of the vector collection; batch embedding with retry and SQLite-backed embedding cache to reduce duplicate API cost.",
        "Retrieval: top-k similarity search with configurable similarity threshold; hybrid reranking coefficient; optional lexical-only demo mode for environments with embedding quota constraints.",
        "Generation: strict system prompting requiring answers to be grounded in retrieved context; fallback string when evidence is insufficient; optional streaming for UX.",
        "Observability: structured logging, request correlation IDs, health (/health) and readiness (/ready) endpoints reflecting dependency state.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    add_heading(doc, "3. Non-functional requirements", 2)
    for line in [
        "Security: secrets via environment variables; CORS configurable per deployment; no secrets committed to source control.",
        "Performance: configurable HTTP timeouts for long-running LLM calls; ingestion throughput suitable for single-node deployment on commodity PaaS (Render-class) free or paid tiers.",
        "Portability: containerised API (Dockerfile); declarative hosting blueprint (e.g. Render YAML); static frontend deployable to Netlify / Cloudflare Pages with build-time API base URL.",
        "Maintainability: modular Python packages (config, ingestion, embeddings, vector store, retriever, RAG chain, API router); automated smoke tests (PowerShell / pytest) for regression detection.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    add_heading(doc, "4. Integration & extensibility", 2)
    add_para(
        doc,
        "The API shall remain stateless with respect to horizontal scaling except for session-scoped "
        "conversation history stored in-process or externalised in a future phase. Integration hooks for "
        "SSO (SAML/OIDC), enterprise secret stores, and observability backends (OpenTelemetry) may be "
        "scoped as Phase 2 deliverables subject to vendor proposal.",
    )

    add_heading(doc, "5. Acceptance & test strategy (technical)", 2)
    add_para(
        doc,
        "Acceptance shall include: successful end-to-end ingest of the reference corpus; grounded responses "
        "on a defined golden-set of questions; verification of citation metadata (section / line or page); "
        "negative tests for out-of-scope queries; load smoke on cold-start hosting if applicable.",
    )

    # --- PART C: Product Owner ---
    doc.add_page_break()
    add_heading(doc, "Part C — Perspective of the Product Owner (Non-Technical)", 1)
    add_heading(doc, "1. What we are buying (in plain language)", 2)
    add_para(
        doc,
        "We want an internal “smart assistant” that behaves like a well-trained analyst who has read a "
        "very long annual report and can answer questions about it quickly. It must not make things up: "
        "if the answer is not in the document, it should say so clearly. When it does answer, it should "
        "point us to where in the document it found the information so we can double-check.",
    )

    add_heading(doc, "2. Who it is for", 2)
    add_para(
        doc,
        "Primary users: strategy, finance, and corporate development colleagues who need reliable "
        "snapshots from public filings for benchmarking and learning—not for automated regulatory filing. "
        "The experience should feel modern and easy: open a web page, (optionally) refresh the index after "
        "an update, type a question, read the answer and citations.",
    )

    add_heading(doc, "3. What “good” looks like to users", 2)
    for line in [
        "Answers are short, clear, and written in business English unless the user asks otherwise.",
        "The assistant introduces itself as tied to OGDCL and explains that it only uses the indexed filing—so expectations are set correctly.",
        "Uploading a new file or re-indexing is a simple button flow, with visible progress or status messages.",
        "If the system is busy or the service is waking up, the user sees a friendly message instead of a cryptic error.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    add_heading(doc, "4. Out of scope (product)", 2)
    add_para(
        doc,
        "The assistant does not provide investment advice, legal opinions, or guaranteed completeness of "
        "the underlying filing. It does not replace Dataroom, EDGAR, or internal document management systems; "
        "it complements them for Q&A-style exploration.",
    )

    add_heading(doc, "5. Roadmap hooks (product backlog)", 2)
    add_para(
        doc,
        "Future enhancements may include: multi-document workspaces; role-based access; Urdu or "
        "bilingual UI; export of Q&A to Word/PDF; analytics on popular questions; integration with "
        "Microsoft 365. Vendors may price these as optional modules.",
    )

    doc.add_page_break()
    add_heading(doc, "Submission instructions", 1)
    add_para(
        doc,
        "Respondents shall submit a single PDF executive proposal (max 25 pages) plus annexes for "
        "technical architecture, pricing, and project plan. Questions shall be submitted via the "
        "programme mailbox by the date published in the procurement calendar. Late submissions may be "
        "excluded at the issuer’s sole discretion.",
    )

    doc.save(out_path)
    print(f"Wrote: {out_path}")


if __name__ == "__main__":
    main()
